"""Document processing service"""

import io
import logging
from typing import List, Tuple
from pathlib import Path

try:
    import pymupdf
except ImportError:
    pymupdf = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from backend.app.core.config import settings
from backend.app.core.exceptions import DocumentProcessingError

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles multi-format document parsing and chunking"""

    SUPPORTED_FORMATS = {".pdf", ".docx", ".txt", ".md"}

    def __init__(self, chunk_size: int = None, chunk_overlap: int = None):
        self.chunk_size = chunk_size or settings.chunk_size
        self.chunk_overlap = chunk_overlap or settings.chunk_overlap

    def process_file(
        self, filename: str, file_content: bytes
    ) -> Tuple[str, List[str], List[dict]]:
        """
        Process uploaded file and return text, chunks, and metadata.

        Returns:
            Tuple of (full_text, chunks, metadata_list)
        """
        file_ext = Path(filename).suffix.lower()

        if file_ext not in self.SUPPORTED_FORMATS:
            raise DocumentProcessingError(
                f"Unsupported file type: {file_ext}. "
                f"Supported: {', '.join(sorted(self.SUPPORTED_FORMATS))}"
            )

        if file_ext == ".pdf":
            return self._process_pdf(filename, file_content)
        elif file_ext == ".docx":
            return self._process_docx(filename, file_content)
        else:
            return self._process_text(filename, file_content)

    def _process_pdf(
        self, filename: str, file_content: bytes
    ) -> Tuple[str, List[str], List[dict]]:
        """Process PDF file with per-page tracking for accurate metadata"""
        if pymupdf is None:
            raise DocumentProcessingError("PDF processing requires PyMuPDF")

        try:
            pdf_bytes = io.BytesIO(file_content)
            doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")

            pages: List[Tuple[int, str]] = []
            full_text = ""
            for page_num in range(len(doc)):
                page_text = doc[page_num].get_text()
                pages.append((page_num + 1, page_text))
                full_text += page_text + "\n"

            doc.close()

            # Chunk per-page to maintain accurate page metadata
            chunks = []
            metadata = []
            for page_num, page_text in pages:
                page_chunks = self._chunk_text(page_text)
                for chunk in page_chunks:
                    chunks.append(chunk)
                    metadata.append({"page": page_num, "source": filename})

            if not chunks:
                chunks = [full_text.strip() or "(empty document)"]
                metadata = [{"page": 1, "source": filename}]

            return full_text, chunks, metadata
        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(f"PDF processing failed: {e}")

    def _process_docx(
        self, filename: str, file_content: bytes
    ) -> Tuple[str, List[str], List[dict]]:
        """Process DOCX file"""
        if DocxDocument is None:
            raise DocumentProcessingError("DOCX processing requires python-docx")

        try:
            docx_bytes = io.BytesIO(file_content)
            doc = DocxDocument(docx_bytes)

            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)

            text = "\n".join(parts)
            chunks = self._chunk_text(text)
            metadata = [{"source": filename} for _ in chunks]

            if not chunks:
                chunks = [text.strip() or "(empty document)"]
                metadata = [{"source": filename}]

            return text, chunks, metadata
        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(f"DOCX processing failed: {e}")

    def _process_text(
        self, filename: str, file_content: bytes
    ) -> Tuple[str, List[str], List[dict]]:
        """Process TXT or MD file"""
        try:
            text = file_content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = file_content.decode("latin-1")
            except Exception as e:
                raise DocumentProcessingError(f"Failed to decode text file: {e}")

        chunks = self._chunk_text(text)
        metadata = [{"source": filename} for _ in chunks]

        if not chunks:
            chunks = [text.strip() or "(empty document)"]
            metadata = [{"source": filename}]

        return text, chunks, metadata

    def _chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks based on word count."""
        words = text.split()
        if not words:
            return []

        words_per_chunk = max(self.chunk_size // 4, 10)
        overlap_words = max(self.chunk_overlap // 4, 5)
        stride = max(words_per_chunk - overlap_words, 1)

        chunks = []
        for i in range(0, len(words), stride):
            chunk = " ".join(words[i : i + words_per_chunk])
            if chunk.strip():
                chunks.append(chunk)
            if i + words_per_chunk >= len(words):
                break

        return chunks
